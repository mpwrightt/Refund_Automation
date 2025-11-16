#!/usr/bin/env python3
"""
Convert WORKFLOW_GUIDE.md to professionally formatted DOCX
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def set_cell_background(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_heading(doc, text, level=1):
    """Add a heading with formatting"""
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(8)
    heading.paragraph_format.space_after = Pt(4)

    if level == 1:
        heading.runs[0].font.size = Pt(16)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        heading.runs[0].font.bold = True
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
    elif level == 2:
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
        heading.runs[0].font.bold = True
        heading.paragraph_format.space_before = Pt(10)
        heading.paragraph_format.space_after = Pt(5)
    elif level == 3:
        heading.runs[0].font.size = Pt(12)
        heading.runs[0].font.color.rgb = RGBColor(51, 102, 153)
        heading.runs[0].font.bold = True
        heading.paragraph_format.space_before = Pt(8)
        heading.paragraph_format.space_after = Pt(4)
    elif level == 4:
        heading.runs[0].font.size = Pt(11)
        heading.runs[0].font.bold = True
        heading.paragraph_format.space_before = Pt(6)
        heading.paragraph_format.space_after = Pt(3)
    return heading

def add_paragraph(doc, text, style=None, bold=False, italic=False):
    """Add a paragraph with compact formatting"""
    p = doc.add_paragraph(text, style=style)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.1

    if bold or italic:
        for run in p.runs:
            run.bold = bold
            run.italic = italic
    return p

def add_code_block(doc, text):
    """Add a compact code block with monospace font"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.4)

    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(51, 51, 51)

    # Add light gray background
    p.style = 'Code'
    return p

def add_table_from_markdown(doc, lines):
    """Parse markdown table and create compact Word table"""
    lines = [l for l in lines if l.strip()]
    if len(lines) < 2:
        return None

    headers = [cell.strip() for cell in lines[0].split('|')[1:-1]]

    rows_data = []
    for line in lines[2:]:
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells:
            rows_data.append(cells)

    if not rows_data:
        return None

    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Compact table spacing
    for row in table.rows:
        row.height = Inches(0.25)
        for cell in row.cells:
            cell.vertical_alignment = 1  # Center
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(1)
                paragraph.paragraph_format.space_after = Pt(1)
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    # Add headers with background color
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_background(header_cells[i], 'D9E2F3')  # Light blue
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    # Add data rows
    for i, row_data in enumerate(rows_data):
        row_cells = table.rows[i + 1].cells
        for j, cell_data in enumerate(row_data):
            if j < len(row_cells):
                row_cells[j].text = cell_data

    return table

def convert_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to professionally formatted Word document"""
    doc = Document()

    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # Modify normal style for compact spacing
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.line_spacing = 1.1

    # Add code style
    try:
        code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Courier New'
        code_style.font.size = Pt(8)
        code_style.paragraph_format.space_before = Pt(3)
        code_style.paragraph_format.space_after = Pt(3)
        code_style.paragraph_format.left_indent = Inches(0.4)
    except:
        pass  # Style already exists

    with open(md_file, 'r') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    in_table = False
    table_lines = []
    code_lines = []
    last_was_heading = False
    last_was_table = False

    while i < len(lines):
        line = lines[i].rstrip()

        # Code blocks
        if line.startswith('```'):
            if in_code_block:
                add_code_block(doc, '\n'.join(code_lines))
                code_lines = []
                in_code_block = False
                last_was_heading = False
                last_was_table = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Tables
        if line.startswith('|') and '|' in line[1:]:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            add_table_from_markdown(doc, table_lines)
            table_lines = []
            in_table = False
            last_was_table = True
            last_was_heading = False

        # Headings
        if line.startswith('# ') and not line.startswith('##'):
            # No extra space before major sections - keep compact
            add_heading(doc, line[2:], level=1)
            last_was_heading = True
            last_was_table = False
        elif line.startswith('## '):
            add_heading(doc, line[3:], level=2)
            last_was_heading = True
            last_was_table = False
        elif line.startswith('### '):
            add_heading(doc, line[4:], level=3)
            last_was_heading = True
            last_was_table = False
        elif line.startswith('#### '):
            add_heading(doc, line[5:], level=4)
            last_was_heading = True
            last_was_table = False

        # Horizontal rules (skip them completely for compact layout)
        elif line.startswith('---'):
            # Skip horizontal rules to save space
            last_was_heading = False
            last_was_table = False

        # Bold text at start
        elif line.startswith('**') and '**' in line[2:]:
            match = re.match(r'\*\*([^*]+)\*\*(.*)', line)
            if match:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(2)
                p.add_run(match.group(1)).bold = True
                p.add_run(match.group(2))
                last_was_heading = False
                last_was_table = False

        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            text = text.replace('**', '').replace('`', '')
            text = text.replace('✓', '✓').replace('✗', '✗').replace('⚠️', '⚠')
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.05
            for run in p.runs:
                run.font.size = Pt(10)
            last_was_heading = False
            last_was_table = False

        # Empty lines (skip most of them for compact layout)
        elif not line.strip():
            # Skip most empty lines to maximize compactness
            last_was_table = False

        # Regular paragraphs
        elif line.strip() and not line.startswith('#'):
            text = line.replace('**', '').replace('`', '')
            text = text.replace('✓', '✓').replace('✗', '✗').replace('⚠️', '⚠')

            # Skip certain metadata lines
            if text.startswith('Document Version:') or text.startswith('Last Updated:') or \
               text.startswith('Audience:') or text.startswith('Classification:'):
                p = doc.add_paragraph(text)
                p.paragraph_format.space_after = Pt(1)
                for run in p.runs:
                    run.font.italic = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(89, 89, 89)
            else:
                p = doc.add_paragraph(text)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(10)

            last_was_heading = False
            last_was_table = False

        i += 1

    # Save document
    doc.save(docx_file)
    print(f"✓ Created {docx_file}")

if __name__ == '__main__':
    convert_markdown_to_docx(
        'docs/WORKFLOW_GUIDE.md',
        'docs/WORKFLOW_GUIDE.docx'
    )
