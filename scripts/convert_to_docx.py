#!/usr/bin/env python3
"""
Convert WORKFLOW_GUIDE.md to formatted DOCX
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def add_heading(doc, text, level=1):
    """Add a heading with formatting"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.size = Pt(24)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    elif level == 2:
        heading.runs[0].font.size = Pt(18)
        heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
    elif level == 3:
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.color.rgb = RGBColor(51, 102, 153)
    return heading

def add_paragraph(doc, text, style=None, bold=False, italic=False):
    """Add a paragraph with optional formatting"""
    p = doc.add_paragraph(text, style=style)
    if bold or italic:
        for run in p.runs:
            run.bold = bold
            run.italic = italic
    return p

def add_code_block(doc, text):
    """Add a code block with monospace font"""
    p = doc.add_paragraph(text)
    p.style = 'Code'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    return p

def add_table_from_markdown(doc, lines):
    """Parse markdown table and create Word table"""
    # Remove empty lines
    lines = [l for l in lines if l.strip()]
    if len(lines) < 2:
        return None

    # Parse header
    headers = [cell.strip() for cell in lines[0].split('|')[1:-1]]

    # Skip separator line
    # Parse data rows
    rows_data = []
    for line in lines[2:]:
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells:
            rows_data.append(cells)

    if not rows_data:
        return None

    # Create table
    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Add data rows
    for i, row_data in enumerate(rows_data):
        row_cells = table.rows[i + 1].cells
        for j, cell_data in enumerate(row_data):
            if j < len(row_cells):
                row_cells[j].text = cell_data

    return table

def convert_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to formatted Word document"""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Add code style
    try:
        code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Courier New'
        code_style.font.size = Pt(9)
    except:
        pass  # Style already exists

    with open(md_file, 'r') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    in_table = False
    table_lines = []
    code_lines = []

    while i < len(lines):
        line = lines[i].rstrip()

        # Code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                add_code_block(doc, '\n'.join(code_lines))
                code_lines = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Tables
        if line.startswith('|') and '|' in line[1:]:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            # End of table
            add_table_from_markdown(doc, table_lines)
            doc.add_paragraph()  # Add space after table
            table_lines = []
            in_table = False

        # Headings
        if line.startswith('# ') and not line.startswith('##'):
            add_heading(doc, line[2:], level=1)
        elif line.startswith('## '):
            add_heading(doc, line[3:], level=2)
        elif line.startswith('### '):
            add_heading(doc, line[4:], level=3)
        elif line.startswith('#### '):
            add_heading(doc, line[5:], level=4)

        # Horizontal rules
        elif line.startswith('---'):
            doc.add_paragraph()

        # Bold text at start (like "**Key Difference:**")
        elif line.startswith('**') and '**' in line[2:]:
            match = re.match(r'\*\*([^*]+)\*\*(.*)', line)
            if match:
                p = doc.add_paragraph()
                p.add_run(match.group(1)).bold = True
                p.add_run(match.group(2))

        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            # Handle bold/italic in bullet points
            text = text.replace('**', '')
            text = text.replace('✓', '✓').replace('✗', '✗').replace('⚠️', '⚠')
            doc.add_paragraph(text, style='List Bullet')

        # Empty lines
        elif not line.strip():
            if i > 0 and lines[i-1].strip():  # Don't add multiple blank lines
                doc.add_paragraph()

        # Regular paragraphs
        elif line.strip() and not line.startswith('#'):
            # Clean up markdown formatting
            text = line.replace('**', '')
            text = text.replace('✓', '✓').replace('✗', '✗').replace('⚠️', '⚠')
            doc.add_paragraph(text)

        i += 1

    # Save document
    doc.save(docx_file)
    print(f"✓ Created {docx_file}")

if __name__ == '__main__':
    convert_markdown_to_docx(
        'docs/WORKFLOW_GUIDE.md',
        'docs/WORKFLOW_GUIDE.docx'
    )
